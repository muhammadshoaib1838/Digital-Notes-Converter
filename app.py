import streamlit as st
import easyocr
import cv2
import numpy as np
from groq import Groq
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from PIL import Image
import io

# --- Page Configuration ---
st.set_page_config(page_title="Document Architect Pro", layout="wide")

# --- CSS for High-Contrast & Button Visibility ---
st.markdown("""
    <style>
    .stApp { background-color: #020617; color: #f8fafc; }
    
    .jaman-tagline {
        background-color: #4c0519;
        color: #ffff00 !important;
        padding: 15px;
        border-radius: 12px;
        text-align: center;
        font-weight: 900;
        font-size: 1.2em;
        border: 1px solid rgba(255,255,255,0.2);
        margin-bottom: 25px;
    }

    div.stButton > button {
        background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%) !important;
        color: #ffffff !important;
        font-weight: bold !important;
        border: none !important;
        width: 100%;
        padding: 10px;
        font-size: 1.1em;
        text-transform: uppercase;
    }
    
    div.stDownloadButton > button {
        background-color: #1e293b !important;
        border: 1px solid #38bdf8 !important;
        color: #38bdf8 !important;
    }

    .signature { color: #c084fc; font-weight: bold; text-align: center; margin-top: 50px; }
    </style>
    """, unsafe_allow_html=True)

# --- API Key Handling ---
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY")

@st.cache_resource
def load_ocr():
    return easyocr.Reader(['en'], gpu=False)

if not GROQ_API_KEY:
    st.error("❌ GROQ_API_KEY missing in Secrets.")
    st.stop()

reader = load_ocr()
client = Groq(api_key=GROQ_API_KEY)

# --- UI Header ---
st.title("🌌 DOCUMENT ARCHITECT PRO")
st.markdown('<div class="jaman-tagline">Transforming handwritten chaos into structured digital gold.</div>', unsafe_allow_html=True)

if "history" not in st.session_state:
    st.session_state.history = []

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("📸 Upload Source")
    uploaded_file = st.file_uploader("Drop image here", type=["jpg", "jpeg", "png"])
    
    if uploaded_file is not None:
        image_preview = Image.open(uploaded_file)
        st.image(image_preview, caption="Preview", use_container_width=True)
        uploaded_file.seek(0)
        
        if st.button("ARCHITECT DOCUMENT ✨"):
            with st.spinner("Styling your document..."):
                try:
                    file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
                    img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

                    if img is None:
                        st.error("Error decoding image.")
                    else:
                        # 1. OCR Stage
                        raw_text = " ".join(reader.readtext(img, detail=0))
                        
                        if not raw_text.strip():
                            st.warning("No text detected in image.")
                        else:
                            # 2. AI Architecture Stage (Colorful Output)
                            prompt = f"""
                            Convert this text into a professional document. 
                            1. Use <h2 style='color:#38bdf8;'>Heading Name</h2> for all headings.
                            2. Use <span style='color:#fbbf24; font-weight:bold;'>Term</span> for important keywords.
                            3. End with '--- FINAL SUMMARY ---'.
                            TEXT: {raw_text}
                            """
                            
                            chat = client.chat.completions.create(
                                messages=[{"role": "user", "content": prompt}],
                                model="llama-3.3-70b-versatile"
                            )
                            full_output = chat.choices[0].message.content
                            
                            # Logic for splitting notes and summary
                            if "--- FINAL SUMMARY ---" in full_output:
                                notes, summary = full_output.split("--- FINAL SUMMARY ---")
                            else:
                                notes, summary = full_output, "Summary included in main text."

                            # 3. Clean Text for Export (Removing HTML tags)
                            clean_text = notes.replace("<h2 style='color:#38bdf8;'>", "").replace("</h2>", "\n").replace("<span style='color:#fbbf24; font-weight:bold;'>", "").replace("</span>", "")

                            # Generate DOCX in Memory
                            doc = Document()
                            doc.add_heading('Architect Pro Export', 0)
                            doc.add_paragraph(clean_text)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            
                            # Generate PDF in Memory
                            pdf_io = io.BytesIO()
                            c = canvas.Canvas(pdf_io, pagesize=letter)
                            c.setFont("Helvetica-Bold", 14)
                            c.drawString(50, 750, "Document Architect Pro Export")
                            c.setFont("Helvetica", 10)
                            y_pos = 720
                            for line in clean_text.split('\n'):
                                if y_pos < 50:
                                    c.showPage()
                                    y_pos = 750
                                c.drawString(50, y_pos, line[:95])
                                y_pos -= 15
                            c.save()

                            # Save results to Session State
                            st.session_state.processed = {
                                "notes": notes.strip(),
                                "summary": summary.strip(),
                                "raw": raw_text,
                                "docx": doc_io.getvalue(),
                                "pdf": pdf_io.getvalue()
                            }
                            st.session_state.history.append(f"Scan {len(st.session_state.history)+1}")

                except Exception as e:
                    st.error(f"Logic Error: {str(e)}")

with col2:
    if "processed" in st.session_state:
        tab1, tab2, tab3 = st.tabs(["✨ Digital Blueprint", "💡 Executive Summary", "📄 Raw OCR"])
        
        with tab1:
            # Render colourful Markdown
            st.markdown(st.session_state.processed["notes"], unsafe_allow_html=True)
            st.divider()
            
            # Download Buttons
            dcol1, dcol2 = st.columns(2)
            with dcol1:
                st.download_button(
                    label="📥 Download DOCX",
                    data=st.session_state.processed["docx"],
                    file_name="Architect_Export.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with dcol2:
                st.download_button(
                    label="📥 Download PDF",
                    data=st.session_state.processed["pdf"],
                    file_name="Architect_Export.pdf",
                    mime="application/pdf"
                )
        
        with tab2:
            st.info(st.session_state.processed["summary"])
            
        with tab3:
            st.text_area("Initial Scan", st.session_state.processed["raw"], height=300)
    else:
        st.info("Upload an image and click 'Architect Document' to begin.")

st.sidebar.subheader("🕒 Session History")
for h in st.session_state.history:
    st.sidebar.write(h)

st.markdown('<div class="signature">DESIGNED & ENGINEERED BY<br>Muhammad Shoaib Nazz</div>', unsafe_allow_html=True)
